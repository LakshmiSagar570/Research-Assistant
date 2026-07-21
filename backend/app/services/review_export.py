"""
FR8: Export generated literature reviews as .docx.

Takes the Markdown produced by review_generator.py and renders it into
a real Word document using python-docx.

Returns raw bytes rather than writing to disk and returning a path.
This is deliberate: on serverless platforms (Vercel), each HTTP request
can be handled by a different function instance with its own ephemeral
/tmp, so a "write file in request A, read it back in request B"
pattern is unreliable there. Building the docx entirely in memory and
returning it in the same request that generated it works identically
in local dev and in serverless production.
"""
import re
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def _add_paragraph_with_bold(doc: Document, text: str, style: str | None = None):
    p = doc.add_paragraph(style=style)
    pos = 0
    for match in _BOLD_RE.finditer(text):
        if match.start() > pos:
            p.add_run(text[pos:match.start()])
        run = p.add_run(match.group(1))
        run.bold = True
        pos = match.end()
    if pos < len(text):
        p.add_run(text[pos:])
    return p


def markdown_to_docx_bytes(markdown_text: str) -> bytes:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for raw_line in markdown_text.split("\n"):
        line = raw_line.rstrip()
        if not line:
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            h = doc.add_heading(line[2:], level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("*") and line.endswith("*") and not line.startswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(line.strip("*"))
            run.italic = True
        elif line.startswith("- "):
            _add_paragraph_with_bold(doc, line[2:], style="List Bullet")
        else:
            _add_paragraph_with_bold(doc, line)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
